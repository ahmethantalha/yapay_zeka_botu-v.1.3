from docx import Document
from typing import BinaryIO, Text, List
import logging
from .base_processor import BaseFileProcessor

class DocxProcessor(BaseFileProcessor):
    def extract_text(self, file: BinaryIO) -> Text:
        """Extract text from DOCX file"""
        try:
            doc = Document(file)
            return "\n\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    def get_metadata(self, file: BinaryIO) -> dict:
        """Get DOCX metadata"""
        try:
            doc = Document(file)
            core_props = doc.core_properties
            return {
                'author': core_props.author or 'Unknown',
                'title': core_props.title or 'Untitled',
                'modified': str(core_props.modified) if core_props.modified else 'Unknown',
                'paragraph_count': len(doc.paragraphs),
                'word_count': sum(len(p.text.split()) for p in doc.paragraphs)
            }
        except Exception as e:
            logging.error(f"Error getting DOCX metadata: {str(e)}")
            return {'paragraph_count': 0, 'word_count': 0}

    def _do_split_file(self, file_path: str, chunk_size: int, method: str) -> List[str]:
        """Split DOCX file into chunks"""
        try:
            self._validate_chunk_size(chunk_size)
            doc = Document(file_path)
            chunks = []
            current_chunk = ""
            
            if method == "page":
                # Approximate page size (500 words per page)
                words_per_page = 500
                current_words = 0
                
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    
                    word_count = len(text.split())
                    current_words += word_count
                    current_chunk += text + "\n\n"
                    
                    if current_words >= words_per_page * chunk_size:
                        chunks.append(current_chunk.strip())
                        current_chunk = ""
                        current_words = 0
            
            else:  # token based
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    
                    current_chunk += text + "\n\n"
                    if self.estimate_tokens(current_chunk) >= chunk_size:
                        chunks.append(current_chunk.strip())
                        current_chunk = ""
            
            # Add remaining content if any
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            
            return chunks

        except Exception as e:
            logging.error(f"Error splitting DOCX file {file_path}: {str(e)}")
            raise