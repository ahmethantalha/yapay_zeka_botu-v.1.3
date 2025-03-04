from dataclasses import dataclass
from datetime import datetime
from typing import Optional, Dict, List
import json
import docx
from fpdf import FPDF
import os

@dataclass
class ProcessingResult:
    original_text: str
    analyzed_text: str
    metadata: dict
    timestamp: datetime
    file_name: str
    analysis_type: str

class ResultManager:
    def __init__(self):
        self.current_result: Optional[ProcessingResult] = None
    
    def save_result(self, result: ProcessingResult, format: str, output_path: str):
        """Save result in specified format"""
        if format == 'json':
            self._save_json(result, output_path)
        elif format == 'txt':
            self._save_txt(result, output_path)
        elif format == 'docx':
            self._save_docx(result, output_path)
        elif format == 'pdf':
            self._save_pdf(result, output_path)
        elif format == 'md':
            self._save_markdown(result, output_path)
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    def _save_json(self, result: ProcessingResult, output_path: str):
        try:
            # Eğer analiz sonucu zaten JSON formatında bir string ise
            try:
                json_data = json.loads(result.analyzed_text)
                
                # JSON dosyasına yaz
                with open(output_path, 'w', encoding='utf-8') as f:
                    json.dump(json_data, f, indent=2, ensure_ascii=False)
                return
            except json.JSONDecodeError:
                pass
                
            # Markdown işaretlerini temizle ve JSON'ı parse et
            text = result.analyzed_text
            # Markdown bloklarını kaldır
            text = text.replace("```json", "").replace("```", "").strip()
            
            # String'i JSON objesine çevir
            json_data = json.loads(text)
            
            # JSON dosyasına yaz
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, indent=2, ensure_ascii=False)
                
        except Exception as e:
            # Hata durumunda orijinal metni sakla
            data = {
                'analyzed_text': result.analyzed_text
            }
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
    
    def _save_txt(self, result: ProcessingResult, output_path: str):
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"Analiz Sonuçları\n")
            f.write(f"===============\n\n")
            f.write(f"Dosya: {result.file_name}\n")
            f.write(f"İşlem Türü: {result.analysis_type}\n")
            f.write(f"Tarih: {result.timestamp}\n\n")
            f.write(result.original_text)
            f.write(f"\n\nSonuç:\n")
            f.write(f"---------\n")
            f.write(result.analyzed_text)
    
    def _save_docx(self, result: ProcessingResult, output_path: str):
        doc = docx.Document()
        
        doc.add_heading('Analiz Sonuçları', 0)
        
        # Add metadata
        doc.add_paragraph(f"Dosya: {result.file_name}")
        doc.add_paragraph(f"Analiz Türü: {result.analysis_type}")
        doc.add_paragraph(f"Tarih: {result.timestamp}")
        
        # Add analysis
        doc.add_heading('Sonuç', level=1)
        doc.add_paragraph(result.analyzed_text)
        
        doc.save(output_path)
    
    def _save_pdf(self, result: ProcessingResult, output_path: str):
        """Save result in PDF format using ReportLab"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Create custom styles for Turkish characters
            styles.add(ParagraphStyle(name='Normal_UTF8', 
                                    fontName='Helvetica',
                                    fontSize=12,
                                    leading=14,
                                    encoding='utf-8'))
            
            styles.add(ParagraphStyle(name='Heading_UTF8', 
                                    fontName='Helvetica-Bold',
                                    fontSize=14,
                                    leading=16,
                                    encoding='utf-8'))
            
            # Create content
            content = []
            
            # Title
            content.append(Paragraph("Analiz Sonuçları", styles['Heading_UTF8']))
            content.append(Spacer(1, 12))
            
            # Metadata
            content.append(Paragraph(f"Dosya: {result.file_name}", styles['Normal_UTF8']))
            content.append(Paragraph(f"Analiz Tipi: {result.analysis_type}", styles['Normal_UTF8']))
            content.append(Paragraph(f"Tarih: {result.timestamp}", styles['Normal_UTF8']))
            content.append(Spacer(1, 24))
            
            # Analysis
            content.append(Paragraph("Analiz", styles['Heading_UTF8']))
            content.append(Spacer(1, 12))
            
            # Split text into paragraphs
            analysis_paras = result.analyzed_text.split('\n')
            for para in analysis_paras:
                if para.strip():  # Skip empty paragraphs
                    content.append(Paragraph(para, styles['Normal_UTF8']))
                    content.append(Spacer(1, 6))
            
            # Build the PDF
            doc.build(content)
            
        except Exception as e:
            raise ValueError(f"PDF oluşturma hatası: {str(e)}")

    def _clean_text(self, text):
        """Remove non-Latin-1 characters from text"""
        return ''.join(c if ord(c) < 256 else '?' for c in text)


    def _save_markdown(self, result: ProcessingResult, output_path: str):
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"# Analysis Results\n\n")
            f.write(f"**File:** {result.file_name}  \n")
            f.write(f"**Analysis Type:** {result.analysis_type}  \n")
            f.write(f"**Date:** {result.timestamp}  \n\n")
            f.write(f"\n\n## Analysis\n\n")
            f.write(result.analyzed_text)

    def combine_results(self, results: List[ProcessingResult], combination_type: str = "sequential") -> ProcessingResult:
        """Birden fazla analiz sonucunu birleştir"""
        if not results:
            raise ValueError("Birleştirilecek sonuç bulunamadı")

        if len(results) == 1:
            return results[0]

        # Birleştirme tipine göre işlem yap
        if combination_type == "sequential":
            return self._combine_sequential(results)
        elif combination_type == "summarize":
            return self._combine_with_summary(results)
        else:
            raise ValueError(f"Geçersiz birleştirme tipi: {combination_type}")

    def _combine_sequential(self, results: List[ProcessingResult]) -> ProcessingResult:
        """Sonuçları sıralı olarak birleştir"""
        all_metadata = {}
        
        # JSON verilerini birleştirmek için
        combined_json = {"soru-cevaplar": []}
        
        for result in results:
            # Metadata'yı birleştir
            for key, value in result.metadata.items():
                if key in all_metadata:
                    if isinstance(all_metadata[key], (int, float)):
                        all_metadata[key] += value
                    elif isinstance(all_metadata[key], list):
                        all_metadata[key].extend(value)
                    else:
                        all_metadata[key] = f"{all_metadata[key]}, {value}"
                else:
                    all_metadata[key] = value
            
            # Eğer analiz sonucu JSON formatındaysa
            try:
                # JSON blok işaretlerini temizle
                text = result.analyzed_text
                text = text.replace("```json", "").replace("```", "").strip()
                
                # Parse JSON
                json_data = json.loads(text)
                
                # Eğer içerisinde "soru-cevaplar" anahtarı varsa, birleştirilmiş JSON'a ekle
                if "soru-cevaplar" in json_data:
                    combined_json["soru-cevaplar"].extend(json_data["soru-cevaplar"])
            except (json.JSONDecodeError, ValueError, AttributeError):
                # JSON parse edilemiyorsa, normal metin olarak ele al
                pass
        
        # Birleştirilmiş JSON'ı string olarak dönüştür
        combined_text = json.dumps(combined_json, indent=2, ensure_ascii=False)
        
        return ProcessingResult(
            original_text="[Birleştirilmiş Metin]",
            analyzed_text=combined_text,
            metadata=all_metadata,
            timestamp=datetime.now(),
            file_name="combined_results.json",
            analysis_type=results[0].analysis_type
        )
    
    def _combine_with_summary(self, results: List[ProcessingResult]) -> ProcessingResult:
        """Sonuçları özetle birleştir"""
        combined_text = "# Analiz Özeti\n\n"
        all_metadata = {}
        
        # Genel bilgiler
        combined_text += f"Toplam Analiz Sayısı: {len(results)}\n"
        combined_text += f"Analiz Tipi: {results[0].analysis_type}\n\n"

        # Her sonuç için özet
        for i, result in enumerate(results, 1):
            combined_text += f"\n## Bölüm {i}\n\n"
            
            # Metni özetle (ilk 200 karakter)
            summary = result.analyzed_text[:200] + "..." if len(result.analyzed_text) > 200 else result.analyzed_text
            combined_text += summary + "\n"

            # Metadata'yı birleştir
            for key, value in result.metadata.items():
                if key in all_metadata:
                    if isinstance(all_metadata[key], (int, float)):
                        all_metadata[key] += value
                    elif isinstance(all_metadata[key], list):
                        all_metadata[key].extend(value)
                    else:
                        all_metadata[key] = f"{all_metadata[key]}, {value}"
                else:
                    all_metadata[key] = value

        # Metadata özeti
        combined_text += "\n# Metadata Özeti\n\n"
        for key, value in all_metadata.items():
            combined_text += f"{key}: {value}\n"

        return ProcessingResult(
            original_text="[Birleştirilmiş Özet]",
            analyzed_text=combined_text,
            metadata=all_metadata,
            timestamp=datetime.now(),
            file_name="combined_summary.txt",
            analysis_type=results[0].analysis_type
        )

    def get_combination_types(self) -> List[Dict[str, str]]:
        """Kullanılabilir birleştirme tiplerini döndür"""
        return [
            {
                "id": "sequential",
                "name": "Sıralı Birleştirme",
                "description": "Tüm sonuçları sıralı olarak birleştirir"
            },
            {
                "id": "summarize",
                "name": "Özetli Birleştirme",
                "description": "Sonuçları özetleyerek birleştirir"
            }
        ]